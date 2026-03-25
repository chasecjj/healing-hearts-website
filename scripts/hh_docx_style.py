"""
Healing Hearts DOCX Style Module
=================================
Reusable style definitions and helper functions for generating
branded Healing Hearts documents from the hub.

Design Standard:
  - Colors from HH brand (teal, coral, cream)
  - Typography: Cormorant Garamond headlines, Arial body (readability)
  - Formatting inspired by structured analysis documents
  - Organization: title page, intro, numbered sections, tables, closing

Usage:
    from hh_docx_style import HHDocument

    doc = HHDocument()
    doc.title_page("Main Title", "Subtitle", "Tagline", "Prepared for X")
    doc.intro_letter("Hi Trisha,", ["paragraph 1", "paragraph 2"])
    doc.section("1. Section Name", "Introduction text")
    doc.subsection("1.1  Subsection Name")
    doc.body("Regular paragraph text")
    doc.labeled("The fix:", "Description of the fix")
    doc.bullet_list(["Item 1", "Item 2", "Item 3"])
    doc.callout("Important note or highlight")
    doc.divider()
    doc.priority_table(headers, rows)
    doc.closing("Sign-off text", "-- PHEDRIS")
    doc.save("output.docx")
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path


# ── Brand Colors ──────────────────────────────────────────────────

TEAL = RGBColor(0x11, 0x91, 0xB1)  # Primary — headings, accents, labels
TEAL_LIGHT = RGBColor(0x7C, 0xC7, 0xD9)  # Light teal — borders, table headers
CORAL = RGBColor(0xB9, 0x6A, 0x5F)  # Secondary — subheadings, highlights
CORAL_LIGHT = RGBColor(0xD4, 0xA0, 0x94)  # Light coral — callout accents
CREAM = RGBColor(0xFA, 0xF4, 0xEA)  # Background — callout fills
DARK = RGBColor(0x17, 0x17, 0x17)  # Body text (never pure black)
GRAY = RGBColor(0x66, 0x66, 0x66)  # Metadata, captions, stage directions
GRAY_LIGHT = RGBColor(0x99, 0x99, 0x99)  # Dates, attribution lines
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

# ── Typography ────────────────────────────────────────────────────

FONT_DISPLAY = "Cormorant Garamond"  # Title page, section headings
FONT_BODY = "Arial"  # Body text, subsections, tables
FONT_ACCENT = "Arial"  # Labels, callouts (bold/italic variants)

# ── Spacing Constants (in EMU) ────────────────────────────────────

SPACE_NONE = Pt(0)
SPACE_XS = Pt(2)
SPACE_SM = Pt(4)
SPACE_MD = Pt(8)
SPACE_LG = Pt(12)
SPACE_XL = Pt(16)
SPACE_XXL = Pt(24)


class HHDocument:
    """Branded Healing Hearts document builder."""

    def __init__(self, margins=1.0):
        self.doc = Document()
        for section in self.doc.sections:
            section.top_margin = Inches(margins)
            section.bottom_margin = Inches(margins)
            section.left_margin = Inches(margins)
            section.right_margin = Inches(margins)

    # ── Primitives ────────────────────────────────────────────────

    def _para(
        self,
        text="",
        font=FONT_BODY,
        size=11,
        color=DARK,
        bold=False,
        italic=False,
        align=None,
        space_before=0,
        space_after=6,
        left_indent=None,
    ):
        """Add a paragraph with full style control."""
        p = self.doc.add_paragraph()
        if text:
            run = p.add_run(text)
            run.font.name = font
            run.font.size = Pt(size)
            run.font.color.rgb = color
            run.bold = bold
            run.italic = italic
        if align is not None:
            p.alignment = align
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after = Pt(space_after)
        if left_indent is not None:
            p.paragraph_format.left_indent = Inches(left_indent)
        return p

    def _add_run(
        self,
        paragraph,
        text,
        font=FONT_BODY,
        size=11,
        color=DARK,
        bold=False,
        italic=False,
    ):
        """Add a styled run to an existing paragraph."""
        run = paragraph.add_run(text)
        run.font.name = font
        run.font.size = Pt(size)
        run.font.color.rgb = color
        run.bold = bold
        run.italic = italic
        return run

    def _set_cell_shading(self, cell, color_hex):
        """Set background color on a table cell."""
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), color_hex)
        shading.set(qn("w:val"), "clear")
        cell._element.get_or_add_tcPr().append(shading)

    def _add_bottom_border(self, paragraph, color_hex="1191B1", size="6", space="4"):
        """Add a colored bottom border to a paragraph (decorative divider)."""
        pPr = paragraph._element.get_or_add_pPr()
        borders = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), size)
        bottom.set(qn("w:space"), space)
        bottom.set(qn("w:color"), color_hex)
        borders.append(bottom)
        pPr.append(borders)

    def _add_left_border(self, paragraph, color_hex="B96A5F", size="12", space="8"):
        """Add a colored left border to a paragraph (callout accent)."""
        pPr = paragraph._element.get_or_add_pPr()
        borders = OxmlElement("w:pBdr")
        left = OxmlElement("w:left")
        left.set(qn("w:val"), "single")
        left.set(qn("w:sz"), size)
        left.set(qn("w:space"), space)
        left.set(qn("w:color"), color_hex)
        borders.append(left)
        pPr.append(borders)

    # ── Title Page ────────────────────────────────────────────────

    def title_page(self, title, subtitle, tagline=None, prepared_for=None):
        """
        Full title page with brand hierarchy.

        Args:
            title: Main document title (e.g., "Spark Challenge")
            subtitle: Secondary title (e.g., "Video Scripts")
            tagline: Optional italic tagline
            prepared_for: Optional attribution line
        """
        # Vertical spacing to center content
        for _ in range(4):
            self._para(space_after=0)

        # Brand mark
        self._para(
            "HEALING HEARTS",
            font=FONT_BODY,
            size=13,
            color=GRAY_LIGHT,
            bold=True,
            align=WD_ALIGN_PARAGRAPH.CENTER,
            space_after=16,
        )

        # Main title
        p = self._para(
            title,
            font=FONT_DISPLAY,
            size=36,
            color=TEAL,
            align=WD_ALIGN_PARAGRAPH.CENTER,
            space_after=4,
        )
        self._add_bottom_border(p, color_hex="B96A5F", size="4", space="8")

        # Subtitle
        self._para(
            subtitle,
            font=FONT_DISPLAY,
            size=22,
            color=CORAL,
            italic=True,
            align=WD_ALIGN_PARAGRAPH.CENTER,
            space_before=8,
            space_after=20,
        )

        # Tagline
        if tagline:
            self._para(
                tagline,
                font=FONT_BODY,
                size=12,
                color=GRAY,
                italic=True,
                align=WD_ALIGN_PARAGRAPH.CENTER,
                space_after=8,
            )

        # Attribution
        if prepared_for:
            self._para(
                prepared_for,
                font=FONT_BODY,
                size=11,
                color=GRAY_LIGHT,
                align=WD_ALIGN_PARAGRAPH.CENTER,
                space_after=0,
            )

        self.doc.add_page_break()

    # ── Intro Letter ──────────────────────────────────────────────

    def intro_letter(self, greeting, paragraphs, signoff="-- PHEDRIS"):
        """
        Warm intro letter (PHEDRIS voice).

        Args:
            greeting: "Hi Trisha," or "Hi Chase,"
            paragraphs: List of body paragraph strings
            signoff: Sign-off line
        """
        self._para(greeting, size=12, bold=True, space_after=10)

        for text in paragraphs:
            self._para(text, size=11, space_after=8)

        self._para(
            signoff, size=11, color=TEAL, italic=True, space_before=12, space_after=0
        )

        self.doc.add_page_break()

    # ── Section Heading (H1) ──────────────────────────────────────

    def section(self, title, intro=None):
        """
        Major section heading with teal color and bottom border.

        Args:
            title: Section title (e.g., "1. Executive Summary")
            intro: Optional introductory paragraph
        """
        p = self._para(
            title,
            font=FONT_DISPLAY,
            size=20,
            color=TEAL,
            bold=True,
            space_before=16,
            space_after=8,
        )
        self._add_bottom_border(p, color_hex="1191B1", size="4", space="4")

        if intro:
            self._para(intro, size=11, space_after=10)

    # ── Subsection Heading (H2) ───────────────────────────────────

    def subsection(self, title):
        """
        Subsection heading in coral (e.g., "1.1  Topic Name").
        """
        self._para(
            title,
            font=FONT_BODY,
            size=14,
            color=CORAL,
            bold=True,
            space_before=14,
            space_after=6,
        )

    # ── Subsubsection (H3) ────────────────────────────────────────

    def subsubsection(self, title):
        """
        Third-level heading in teal, smaller.
        """
        self._para(
            title,
            font=FONT_BODY,
            size=12,
            color=TEAL,
            bold=True,
            space_before=10,
            space_after=4,
        )

    # ── Body Text ─────────────────────────────────────────────────

    def body(self, text, space_after=10):
        """Standard body paragraph."""
        self._para(text, size=11, space_after=space_after)

    # ── Labeled Paragraph ─────────────────────────────────────────

    def labeled(self, label, text):
        """
        Paragraph with a bold teal label prefix.
        e.g., labeled("The fix:", "Add a feedback loop...")
        """
        p = self._para(space_after=10)
        self._add_run(p, label + " ", size=11, color=TEAL, bold=True)
        self._add_run(p, text, size=11, color=DARK)

    # ── Bullet List ───────────────────────────────────────────────

    def bullet_list(self, items, indent=0.4):
        """Indented bullet list with teal bullets."""
        for item in items:
            p = self._para(space_after=5, left_indent=indent)
            self._add_run(p, "\u2022  ", size=11, color=TEAL, bold=True)
            self._add_run(p, item, size=11, color=DARK)

    # ── Numbered List ─────────────────────────────────────────────

    def numbered_list(self, items, indent=0.4):
        """Numbered list with teal numbers."""
        for i, item in enumerate(items, 1):
            p = self._para(space_after=5, left_indent=indent)
            self._add_run(p, f"{i}.  ", size=11, color=TEAL, bold=True)
            self._add_run(p, item, size=11, color=DARK)

    # ── Callout Box ───────────────────────────────────────────────

    def callout(self, text, label=None):
        """
        Highlighted callout with coral left border.
        Optional label (e.g., "Note:", "Important:").
        """
        p = self._para(space_before=8, space_after=8, left_indent=0.3)
        self._add_left_border(p, color_hex="B96A5F", size="18", space="10")

        if label:
            self._add_run(p, label + " ", size=11, color=CORAL, bold=True)

        self._add_run(p, text, size=11, color=DARK, italic=True)

    # ── Divider ───────────────────────────────────────────────────

    def divider(self):
        """Thin teal horizontal line between content sections."""
        p = self._para(space_before=8, space_after=8)
        self._add_bottom_border(p, color_hex="7CC7D9", size="2", space="0")

    # ── Script Block ──────────────────────────────────────────────

    def script_section(self, section_name):
        """Section header for video/content scripts (teal, uppercase)."""
        self._para(
            section_name.upper(),
            font=FONT_BODY,
            size=11,
            color=TEAL,
            bold=True,
            space_before=10,
            space_after=4,
        )

    def script_line(self, text):
        """Spoken script line (body text, tight spacing)."""
        self._para(text, size=11, space_after=2)

    def stage_direction(self, text):
        """Stage direction or emotion cue (gray italic)."""
        self._para(text, size=10, color=GRAY, italic=True, space_after=2)

    def production_note(self, text):
        """Production note (coral italic, indented)."""
        p = self._para(
            text, size=10, color=CORAL, italic=True, space_after=4, left_indent=0.3
        )

    def speaker_label(self, name):
        """Speaker label for multi-person scripts (e.g., TRISHA:)."""
        self._para(
            name + ":",
            font=FONT_BODY,
            size=10,
            color=TEAL,
            bold=True,
            space_before=8,
            space_after=4,
        )

    # ── Review Prompt ─────────────────────────────────────────────

    def review_prompt(self, text):
        """Reviewer guidance (teal label + gray italic text)."""
        p = self._para(space_after=10)
        self._add_run(p, "For your review: ", size=10, color=TEAL, bold=True)
        self._add_run(p, text, size=10, color=GRAY, italic=True)

    # ── Table ─────────────────────────────────────────────────────

    def table(self, headers, rows, col_widths=None):
        """
        Branded table with teal header row.

        Args:
            headers: List of column header strings
            rows: List of lists (each inner list = one row)
            col_widths: Optional list of column widths in inches
        """
        tbl = self.doc.add_table(rows=1, cols=len(headers))
        tbl.style = "Table Grid"
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Header row
        for i, header in enumerate(headers):
            cell = tbl.rows[0].cells[i]
            self._set_cell_shading(cell, "1191B1")
            p = cell.paragraphs[0]
            run = p.add_run(header)
            run.font.name = FONT_BODY
            run.font.size = Pt(10)
            run.font.color.rgb = WHITE
            run.bold = True

        # Data rows
        for row_data in rows:
            row = tbl.add_row()
            for i, val in enumerate(row_data):
                cell = row.cells[i]
                p = cell.paragraphs[0]
                run = p.add_run(str(val))
                run.font.name = FONT_BODY
                run.font.size = Pt(10)
                run.font.color.rgb = DARK

        # Column widths
        if col_widths:
            for i, width in enumerate(col_widths):
                for row in tbl.rows:
                    row.cells[i].width = Inches(width)

        # Add spacing after table
        self._para(space_after=10)

    # ── Metadata Line ─────────────────────────────────────────────

    def metadata(self, text):
        """Metadata line (small, gray, italic). For dates, attribution, stats."""
        self._para(text, size=10, color=GRAY, italic=True, space_after=10)

    # ── Page Break ────────────────────────────────────────────────

    def page_break(self):
        self.doc.add_page_break()

    # ── Closing ───────────────────────────────────────────────────

    def closing(self, text, signoff="-- PHEDRIS"):
        """Warm closing paragraph with sign-off."""
        self.doc.add_page_break()
        self.section("What Happens Next")
        self._para(text, size=11, space_after=12)
        self._para(signoff, size=11, color=TEAL, italic=True, space_after=0)

    def closing_with_steps(self, steps, final_note=None, signoff="-- PHEDRIS"):
        """Closing page with numbered next steps."""
        self.doc.add_page_break()
        self.section("What Happens Next")
        self.numbered_list(steps)

        if final_note:
            self._para("", space_after=8)
            self._para(final_note, size=11, space_after=12)

        self._para(
            signoff, size=11, color=TEAL, italic=True, space_before=12, space_after=0
        )

    # ── Save ──────────────────────────────────────────────────────

    def save(self, path):
        """Save the document. Creates parent directories if needed."""
        path = Path(path)
        path.parent.mkdir(parents=True, exist_ok=True)
        self.doc.save(str(path))
        return path
