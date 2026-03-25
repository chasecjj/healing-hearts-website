"""
Build a warm, Trisha-friendly review document for the 7 Spark Challenge video scripts.
Written as PHEDRIS would write to Trisha -- no code, no file paths, no jargon.
"""

import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Paths
SCRIPTS_DIR = Path(
    r"C:\Users\chase\Documents\HealingHeartsWebsite\docs\video-scripts\spark-challenge"
)
OUTPUT = Path(
    r"C:\Users\chase\Documents\Healing Hearts\Spark Challenge Video Scripts - Trisha Review.docx"
)

# Design tokens (Healing Hearts brand-adjacent)
TEAL = RGBColor(0x11, 0x91, 0xB1)
CORAL = RGBColor(0xB9, 0x6A, 0x5F)
DARK = RGBColor(0x17, 0x17, 0x17)
GRAY = RGBColor(0x66, 0x66, 0x66)


def strip_frontmatter(text):
    """Remove YAML frontmatter from markdown."""
    if text.startswith("---"):
        end = text.find("---", 3)
        if end != -1:
            return text[end + 3 :].strip()
    return text.strip()


def parse_script(filepath):
    """Parse a script markdown file into sections."""
    raw = filepath.read_text(encoding="utf-8")
    body = strip_frontmatter(raw)

    # Extract frontmatter values
    fm = {}
    if raw.startswith("---"):
        fm_text = raw[3 : raw.find("---", 3)]
        for line in fm_text.strip().split("\n"):
            if ":" in line and not line.startswith("  -"):
                key, val = line.split(":", 1)
                fm[key.strip()] = val.strip().strip('"')

    # Clean up markdown artifacts
    # Remove H1 headers like "# Day 3 -- The 2-Minute Check-In"
    body = re.sub(r"^# Day \d+ -- .+$", "", body, flags=re.MULTILINE)
    # Remove "## Beat Sheet"
    body = re.sub(r"^## Beat Sheet.*$", "", body, flags=re.MULTILINE)
    # Remove standalone "---" dividers
    body = re.sub(r"^---+\s*$", "", body, flags=re.MULTILINE)
    # Remove "## Teleprompter Notes" section and everything after
    body = re.sub(r"## Teleprompter Notes.*", "", body, flags=re.DOTALL)

    sections = []
    current_section = None
    current_lines = []

    for line in body.split("\n"):
        # Match section headers like "## HOOK (10-15s)" or "### HOOK"
        header_match = re.match(r"^#{2,3}\s+(.+?)(?:\s*\(.*\))?\s*(?:--\s*.*)?$", line)
        if header_match:
            if current_section:
                sections.append((current_section, "\n".join(current_lines).strip()))
            current_section = header_match.group(1).strip()
            # Clean up section names
            current_section = re.sub(r"\s*--\s*.*$", "", current_section)
            current_lines = []
        else:
            current_lines.append(line)

    if current_section:
        sections.append((current_section, "\n".join(current_lines).strip()))

    return fm, sections


def add_styled_paragraph(
    doc,
    text,
    font_name="Plus Jakarta Sans",
    size=11,
    color=DARK,
    bold=False,
    italic=False,
    alignment=None,
    space_after=6,
):
    """Add a styled paragraph."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    run.italic = italic
    if alignment:
        p.alignment = alignment
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    return p


def add_script_line(doc, line):
    """Add a single script line with appropriate formatting."""
    line = line.strip()
    if not line:
        return

    # Production notes (blockquotes)
    if line.startswith(">"):
        text = line.lstrip("> ").strip()
        if text.startswith("PRODUCTION NOTE") or text.startswith("PRODUCTION NOTES"):
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.font.name = "Plus Jakarta Sans"
            run.font.size = Pt(10)
            run.font.color.rgb = CORAL
            run.italic = True
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.left_indent = Inches(0.3)
            return
        # Other production note lines
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = "Plus Jakarta Sans"
        run.font.size = Pt(10)
        run.font.color.rgb = CORAL
        run.italic = True
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.left_indent = Inches(0.3)
        return

    # Stage directions in brackets
    if line.startswith("[") and line.endswith("]"):
        p = doc.add_paragraph()
        run = p.add_run(line)
        run.font.name = "Plus Jakarta Sans"
        run.font.size = Pt(10)
        run.font.color.rgb = GRAY
        run.italic = True
        p.paragraph_format.space_after = Pt(2)
        return

    # Speaker labels (TRISHA:, JEFF:, etc.)
    if re.match(r"^(TRISHA|JEFF|TRISHA \+ JEFF|TRISHA AND JEFF)\s*:", line):
        p = doc.add_paragraph()
        run = p.add_run(line)
        run.font.name = "Outfit"
        run.font.size = Pt(10)
        run.font.color.rgb = TEAL
        run.bold = True
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.space_before = Pt(8)
        return

    # Regular spoken text
    p = doc.add_paragraph()
    run = p.add_run(line)
    run.font.name = "Plus Jakarta Sans"
    run.font.size = Pt(11)
    run.font.color.rgb = DARK
    p.paragraph_format.space_after = Pt(2)


def build_document():
    doc = Document()

    # Set default margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)

    # --- TITLE PAGE ---
    for _ in range(4):
        doc.add_paragraph()

    add_styled_paragraph(
        doc,
        "Spark Challenge",
        font_name="Cormorant Garamond",
        size=36,
        color=TEAL,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=4,
    )
    add_styled_paragraph(
        doc,
        "Companion Video Scripts",
        font_name="Cormorant Garamond",
        size=24,
        color=CORAL,
        italic=True,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=24,
    )
    add_styled_paragraph(
        doc,
        "For Trisha's Review",
        font_name="Plus Jakarta Sans",
        size=14,
        color=GRAY,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=8,
    )
    add_styled_paragraph(
        doc,
        "Prepared by PHEDRIS  |  March 2026",
        font_name="Plus Jakarta Sans",
        size=11,
        color=GRAY,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=0,
    )

    doc.add_page_break()

    # --- INTRO LETTER ---
    add_styled_paragraph(
        doc,
        "Hi Trisha,",
        font_name="Plus Jakarta Sans",
        size=12,
        color=DARK,
        bold=True,
        space_after=12,
    )

    intro_paragraphs = [
        "These are the companion video scripts for the 7-Day Spark Challenge. Each one is designed to be about 90 seconds to 2 minutes -- short enough to hold attention, long enough to land emotionally.",
        'Your job in these videos is simple: look at the camera, tell the story, and make people feel something. The daily emails handle the exercise instructions. These videos are the emotional anchor -- the moment where someone watching thinks, "She gets it. She gets me."',
        "I pulled as much as I could from your real coaching sessions -- the stories you've told, the phrases you actually use, the way you teach. Some scripts are heavily sourced from your sessions (80%+), and a few needed more creative bridging where the transcripts were thin. I've flagged those honestly.",
        "Here's what I need from you:",
    ]

    for para in intro_paragraphs:
        add_styled_paragraph(doc, para, size=11, space_after=8)

    # Review checklist
    review_items = [
        "Read each script out loud. Does it sound like you? If a phrase feels off, cross it out and write what you'd actually say.",
        "The stories -- are these the right ones? If you have a better story for any day, swap it in. Your real stories will always beat anything I assembled.",
        'Day 6 especially -- I couldn\'t find your "moment I knew" story about Jeff in the transcripts. I used your coaching exercise and the "old people on the couch" line instead. During the dry run, I\'d love for you to just tell that story in your own words and we\'ll use that.',
        "Day 7 is partially unscripted -- you'll ask Jeff the question live on camera. His answer is real, not rehearsed. Your setup and closing are on the teleprompter. The middle is just... you two being you.",
        "Mark anything that needs Jeff's input -- especially Day 4 where he explains the Critter Brain science.",
    ]

    for item in review_items:
        p = doc.add_paragraph()
        # Add bullet
        run = p.add_run("  \u2022  ")
        run.font.name = "Plus Jakarta Sans"
        run.font.size = Pt(11)
        run.font.color.rgb = TEAL
        run.bold = True
        # Add text
        run2 = p.add_run(item)
        run2.font.name = "Plus Jakarta Sans"
        run2.font.size = Pt(11)
        run2.font.color.rgb = DARK
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.left_indent = Inches(0.2)

    add_styled_paragraph(doc, "", size=6, space_after=6)

    add_styled_paragraph(
        doc,
        "Take your time with these. Cross things out, write in the margins, circle what you love. Nothing is final until you say it is.",
        size=11,
        space_after=12,
    )
    add_styled_paragraph(
        doc, "-- PHEDRIS", size=11, color=TEAL, italic=True, space_after=0
    )

    doc.add_page_break()

    # --- QUICK REFERENCE ---
    add_styled_paragraph(
        doc,
        "Quick Reference",
        font_name="Cormorant Garamond",
        size=22,
        color=TEAL,
        space_after=12,
    )

    day_summaries = [
        (
            "Day 1",
            'The "I Noticed" Text',
            "You solo",
            "Jeff's lunch visits + client gratitude letter",
        ),
        (
            "Day 2",
            "The Specific Spark Compliment",
            "You solo",
            '"Checking a box" client story + specific praise',
        ),
        ("Day 3", "The 2-Minute Check-In", "You solo", "The smartphone boy story"),
        (
            "Day 4",
            "The Pause Experiment",
            "You + Jeff",
            "Critter Brain / CEO Brain -- Jeff explains the science",
        ),
        (
            "Day 5",
            "The Gratitude Text",
            "You solo",
            "Jeff lunch surprise + David's flowers",
        ),
        (
            "Day 6",
            "The Memory Lane Moment",
            "You solo",
            'Coaching memory exercise + "old people on the couch"',
        ),
        (
            "Day 7",
            "The Spark Conversation",
            "You + Jeff",
            "You ask Jeff the question live -- unscripted",
        ),
    ]

    # Table
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"

    # Header row
    headers = ["Day", "Title", "Who's On Camera", "Key Story"]
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        p = cell.paragraphs[0]
        run = p.add_run(header)
        run.font.name = "Plus Jakarta Sans"
        run.font.size = Pt(10)
        run.font.color.rgb = TEAL
        run.bold = True

    for day, title, talent, story in day_summaries:
        row = table.add_row()
        vals = [day, title, talent, story]
        for i, val in enumerate(vals):
            cell = row.cells[i]
            p = cell.paragraphs[0]
            run = p.add_run(val)
            run.font.name = "Plus Jakarta Sans"
            run.font.size = Pt(10)
            run.font.color.rgb = DARK

    doc.add_page_break()

    # --- INDIVIDUAL SCRIPTS ---
    script_files = sorted(SCRIPTS_DIR.glob("day-*.md"))

    day_titles = {
        1: 'The "I Noticed" Text',
        2: "The Specific Spark Compliment",
        3: "The 2-Minute Check-In",
        4: "The Pause Experiment",
        5: "The Gratitude Text",
        6: "The Memory Lane Moment",
        7: "The Spark Conversation",
    }

    day_review_prompts = {
        1: "Does the Jeff lunch story feel right? Is that how you'd tell it? And the client gratitude letter -- is that a moment you remember well enough to own on camera?",
        2: 'The "checking a box" client story -- does that land for you? And the cooking detail at the end -- would you keep that or swap in something stronger?',
        3: "The smartphone story is the heart of this one. You've told it before in sessions -- does this version feel natural? Would you tell it differently?",
        4: "This is the one where Jeff steps in. Does the handoff feel smooth? And Jeff's section -- is the vagus nerve / \"reset switch\" language something he'd be comfortable saying on camera?",
        5: "Two stories here -- your lunch surprise for Jeff, and David's flowers. Do both feel authentic? Is there a gratitude moment you'd rather use?",
        6: 'This is the script I\'m least sure about. I couldn\'t find your early love story in the transcripts, so I used the coaching exercise and the "old people on the couch" line. During the dry run, would you just tell your real "moment I knew" story? Whatever you say will replace what\'s here.',
        7: "Your setup and closing are scripted. The middle -- asking Jeff the question -- is real. No rehearsal. Are you comfortable with that? And does the course mention at the end feel warm enough, or too salesy?",
    }

    for filepath in script_files:
        fm, sections = parse_script(filepath)

        # Get day number
        day_num = int(fm.get("day", "0"))
        title = day_titles.get(day_num, fm.get("title", ""))

        # Day header
        add_styled_paragraph(
            doc,
            f"Day {day_num}",
            font_name="Cormorant Garamond",
            size=28,
            color=TEAL,
            space_after=2,
        )
        add_styled_paragraph(
            doc,
            title,
            font_name="Cormorant Garamond",
            size=18,
            color=CORAL,
            italic=True,
            space_after=4,
        )

        # Talent + duration
        talent = fm.get("talent", "")
        word_count = fm.get("word_count", "")
        coverage = fm.get("transcript_coverage", "")

        meta_text = (
            f"{talent}  |  ~{word_count} words  |  {coverage} from your real sessions"
        )
        add_styled_paragraph(
            doc, meta_text, size=10, color=GRAY, italic=True, space_after=12
        )

        # Review prompt for this day
        if day_num in day_review_prompts:
            p = doc.add_paragraph()
            run = p.add_run("For your review: ")
            run.font.name = "Plus Jakarta Sans"
            run.font.size = Pt(10)
            run.font.color.rgb = TEAL
            run.bold = True
            run2 = p.add_run(day_review_prompts[day_num])
            run2.font.name = "Plus Jakarta Sans"
            run2.font.size = Pt(10)
            run2.font.color.rgb = GRAY
            run2.italic = True
            p.paragraph_format.space_after = Pt(12)

        # Script sections
        for section_name, section_text in sections:
            # Section header
            add_styled_paragraph(
                doc,
                section_name.upper(),
                font_name="Outfit",
                size=11,
                color=TEAL,
                bold=True,
                space_after=4,
            )

            # Script lines
            for line in section_text.split("\n"):
                add_script_line(doc, line)

            # Small spacer between sections
            add_styled_paragraph(doc, "", size=4, space_after=8)

        # Page break between days (except after last)
        if day_num < 7:
            doc.add_page_break()

    # --- CLOSING PAGE ---
    doc.add_page_break()

    add_styled_paragraph(
        doc,
        "What Happens Next",
        font_name="Cormorant Garamond",
        size=22,
        color=TEAL,
        space_after=12,
    )

    next_steps = [
        "You review these scripts -- mark up anything that doesn't sound like you.",
        "We do a teleprompter dry run with Day 1 to set your scroll speed and catch anything that feels awkward when spoken.",
        'Day 6: you tell your real "moment I knew" story on the spot. We\'ll use your words.',
        "Filming session 1: Days 1, 2, 3, 5, 6 -- just you and the camera.",
        "Filming session 2: Days 4 and 7 -- Jeff joins you.",
        "Videos go up before the expo on April 10.",
    ]

    for i, step in enumerate(next_steps, 1):
        p = doc.add_paragraph()
        run = p.add_run(f"  {i}.  ")
        run.font.name = "Plus Jakarta Sans"
        run.font.size = Pt(11)
        run.font.color.rgb = TEAL
        run.bold = True
        run2 = p.add_run(step)
        run2.font.name = "Plus Jakarta Sans"
        run2.font.size = Pt(11)
        run2.font.color.rgb = DARK
        p.paragraph_format.space_after = Pt(6)

    add_styled_paragraph(doc, "", size=8, space_after=12)

    add_styled_paragraph(
        doc,
        "Nothing in this document is final until you say it is. These are your words, your stories, your course. I just helped organize them.",
        size=11,
        space_after=12,
    )
    add_styled_paragraph(
        doc,
        "You're going to be great on camera, Trisha.",
        size=12,
        color=TEAL,
        italic=True,
        bold=True,
        space_after=4,
    )
    add_styled_paragraph(
        doc, "-- PHEDRIS", size=11, color=TEAL, italic=True, space_after=0
    )

    # Save
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUTPUT))
    print(f"Saved to: {OUTPUT}")


if __name__ == "__main__":
    build_document()
